"""
Dosya İşlemleri Modülü — Dinamik & Esnek v2.

Tüm metotlar minimum veri ile çalışır:
  - Tam yol verilmezse dosya adından otomatik arama yapar
  - Parametre isimleri esnektir (path, file_path, name hepsi kabul edilir)
  - Dizin yolları da otomatik çözümlenir

Bu modül BaseModule'den türer ve otomatik olarak sisteme kaydolur.
"""

import os
import shutil
import time
import json
from pathlib import Path
from datetime import datetime
from typing import Optional, Union

import sys

from core.BaseModule import BaseModule, ModuleMetadata, Tool, ToolParam
from indexing.FileQueryEngine import FileQueryEngine

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


class FileOperationsModule(BaseModule):
    """Dosya sistemi işlemleri modülü — Dinamik parametre çözümleme destekli."""

    def __init__(self, config_path: str = "config.json"):
        self.config_path = config_path
        self.query_engine = FileQueryEngine(config_path)

    # ================================================================
    # METADATA & TOOLS
    # ================================================================

    @property
    def metadata(self) -> ModuleMetadata:
        return ModuleMetadata(
            name="file_operations",
            description="Dosya arama, okuma, kopyalama, taşıma, silme, oluşturma ve dizin işlemleri",
            version="2.0.0",
        )

    @property
    def tools(self) -> list[Tool]:
        return [
            # --- ARAMA ---
            Tool(
                name="search_files",
                description="Dosya ara. İsim, uzantı, dizin, boyut filtrelerini destekler.",
                params=[
                    ToolParam(
                        "name", "Dosya adında aranacak metin", "string", required=False
                    ),
                    ToolParam(
                        "extension",
                        "Dosya uzantısı (.py, .txt vs.)",
                        "string",
                        required=False,
                    ),
                    ToolParam(
                        "directory",
                        "Aranacak dizin adı veya yolu",
                        "string",
                        required=False,
                    ),
                    ToolParam(
                        "min_size_mb", "Minimum boyut (MB)", "number", required=False
                    ),
                    ToolParam(
                        "max_size_mb", "Maximum boyut (MB)", "number", required=False
                    ),
                    ToolParam(
                        "limit",
                        "Maksimum sonuç sayısı",
                        "number",
                        required=False,
                        default=20,
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="search_directories",
                description="Dizin/klasör ara.",
                params=[
                    ToolParam(
                        "name", "Dizin adında aranacak metin", "string", required=False
                    ),
                    ToolParam(
                        "parent_path", "Üst dizin yolu", "string", required=False
                    ),
                    ToolParam(
                        "limit",
                        "Maksimum sonuç sayısı",
                        "number",
                        required=False,
                        default=20,
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="find_recently_modified",
                description="Son X gün içinde değiştirilen dosyaları bul.",
                params=[
                    ToolParam(
                        "days", "Kaç gün geriye bakılacak", "number", required=True
                    ),
                    ToolParam("extension", "Uzantı filtresi", "string", required=False),
                    ToolParam("directory", "Dizin filtresi", "string", required=False),
                ],
                danger_level="safe",
            ),
            Tool(
                name="list_directory_contents",
                description="Bir dizinin içeriğini listele (dosyalar ve alt dizinler).",
                params=[
                    ToolParam(
                        "path",
                        "Dizin yolu veya ismi (tam yol veya sadece isim)",
                        "string",
                        required=True,
                    ),
                ],
                danger_level="safe",
            ),
            # --- OKUMA ---
            Tool(
                name="read_file",
                description="Dosyanın içeriğini oku ve göster. Tam yol veya sadece dosya adı verilebilir.",
                params=[
                    ToolParam(
                        "path",
                        "Dosya yolu veya adı (ör: test.txt veya /home/user/test.txt)",
                        "string",
                        required=True,
                    ),
                    ToolParam(
                        "max_lines",
                        "Maksimum satır sayısı",
                        "number",
                        required=False,
                        default=50,
                    ),
                ],
                danger_level="safe",
            ),
            # --- YAZMA ---
            Tool(
                name="write_file",
                description="Dosyaya metin yaz. Dosya varsa üzerine yazar, yoksa oluşturur.",
                params=[
                    ToolParam("path", "Dosya yolu veya adı", "string", required=True),
                    ToolParam("content", "Yazılacak içerik", "string", required=True),
                ],
                danger_level="safe",
            ),
            # --- METİN ANALİZ ---
            Tool(
                name="find_in_text",
                description="Metin içinde URL, IP adresi, e-posta, domain veya özel kalıp arar. Pattern: url, ip, email, domain veya regex.",
                params=[
                    ToolParam(
                        "text",
                        "Aranacak metin (veya önceki adımın çıktısı)",
                        "string",
                        required=True,
                    ),
                    ToolParam(
                        "pattern",
                        "Aranacak kalıp tipi: url, ip, email, domain veya regex ifadesi",
                        "string",
                        required=True,
                    ),
                ],
                danger_level="safe",
            ),
            # --- KOPYALAMA / TAŞIMA ---
            Tool(
                name="copy_file",
                description="Dosyayı bir konumdan başka bir konuma kopyala.",
                params=[
                    ToolParam(
                        "source", "Kaynak dosya yolu veya adı", "string", required=True
                    ),
                    ToolParam(
                        "destination",
                        "Hedef yol (dizin veya tam dosya yolu)",
                        "string",
                        required=True,
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="move_file",
                description="Dosyayı taşı (kes-yapıştır).",
                params=[
                    ToolParam(
                        "source", "Kaynak dosya yolu veya adı", "string", required=True
                    ),
                    ToolParam("destination", "Hedef yol", "string", required=True),
                ],
                danger_level="confirm",
            ),
            # --- OLUŞTURMA ---
            Tool(
                name="create_file",
                description="Yeni bir boş dosya oluştur.",
                params=[
                    ToolParam(
                        "path", "Oluşturulacak dosyanın yolu", "string", required=True
                    ),
                    ToolParam(
                        "content",
                        "Dosya içeriği (opsiyonel)",
                        "string",
                        required=False,
                        default="",
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="create_directory",
                description="Yeni bir dizin/klasör oluştur.",
                params=[
                    ToolParam(
                        "path", "Oluşturulacak dizinin yolu", "string", required=True
                    ),
                ],
                danger_level="safe",
            ),
            # --- SİLME ---
            Tool(
                name="delete_file",
                description="Dosya sil. DİKKAT: Geri alınamaz!",
                params=[
                    ToolParam(
                        "path", "Silinecek dosya yolu veya adı", "string", required=True
                    ),
                ],
                danger_level="dangerous",
            ),
            Tool(
                name="delete_directory",
                description="Dizin ve içeriğini sil. DİKKAT: Geri alınamaz!",
                params=[
                    ToolParam(
                        "path", "Silinecek dizin yolu veya adı", "string", required=True
                    ),
                ],
                danger_level="dangerous",
            ),
            # --- YENİDEN ADLANDIRMA ---
            Tool(
                name="rename",
                description="Dosya veya dizini yeniden adlandır.",
                params=[
                    ToolParam(
                        "path", "Dosya/dizin yolu veya adı", "string", required=True
                    ),
                    ToolParam(
                        "new_name",
                        "Yeni ad (sadece isim, yol değil)",
                        "string",
                        required=True,
                    ),
                ],
                danger_level="confirm",
            ),
            # --- BİLGİ ---
            Tool(
                name="get_file_info",
                description="Dosya hakkında detaylı bilgi getir (boyut, tarih, yol).",
                params=[
                    ToolParam("path", "Dosya yolu veya adı", "string", required=True),
                ],
                danger_level="safe",
            ),
            Tool(
                name="find_directory_size",
                description="Dizinin toplam boyutunu hesapla.",
                params=[
                    ToolParam("path", "Dizin yolu veya adı", "string", required=True),
                ],
                danger_level="safe",
            ),
            # --- İÇERİK İŞLEMLERİ ---
            Tool(
                name="search_in_files",
                description="Dosyaların içinde metin veya regex ara.",
                params=[
                    ToolParam(
                        "query", "Aranacak metin veya regex", "string", required=True
                    ),
                    ToolParam("directory", "Aranacak dizin", "string", required=False),
                    ToolParam("extension", "Uzantı filtresi", "string", required=False),
                    ToolParam(
                        "is_regex",
                        "Regex kullan",
                        "boolean",
                        required=False,
                        default=False,
                    ),
                    ToolParam(
                        "limit", "Maks sonuç", "number", required=False, default=20
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="append_to_file",
                description="Dosyanın sonuna içerik ekle (mevcut içeriği silmez).",
                params=[
                    ToolParam("path", "Dosya yolu veya adı", "string", required=True),
                    ToolParam("content", "Eklenecek içerik", "string", required=True),
                ],
                danger_level="safe",
            ),
            Tool(
                name="replace_text",
                description="Dosya içindeki metni bul ve değiştir.",
                params=[
                    ToolParam("path", "Dosya yolu veya adı", "string", required=True),
                    ToolParam(
                        "old_text", "Değiştirilecek metin", "string", required=True
                    ),
                    ToolParam("new_text", "Yeni metin", "string", required=True),
                    ToolParam(
                        "is_regex",
                        "Regex kullan",
                        "boolean",
                        required=False,
                        default=False,
                    ),
                ],
                danger_level="confirm",
            ),
            Tool(
                name="count_lines_words",
                description="Dosyadaki satır, kelime ve karakter sayısını göster.",
                params=[
                    ToolParam("path", "Dosya yolu veya adı", "string", required=True),
                ],
                danger_level="safe",
            ),
            # --- BÜYÜK/KÜÇÜK DOSYALAR ---
            Tool(
                name="find_largest_files",
                description="En büyük dosyaları bul. Boyut filtresi opsiyonel.",
                params=[
                    ToolParam(
                        "min_size_mb",
                        "Minimum boyut MB cinsinden",
                        "number",
                        required=False,
                    ),
                    ToolParam("extension", "Uzantı filtresi", "string", required=False),
                    ToolParam("directory", "Dizin filtresi", "string", required=False),
                    ToolParam(
                        "limit", "Kaç dosya", "number", required=False, default=10
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="find_small_files",
                description="Belirli boyutun altındaki dosyaları bul.",
                params=[
                    ToolParam(
                        "max_size_mb",
                        "Maksimum boyut MB cinsinden",
                        "number",
                        required=False,
                        default=1.0,
                    ),
                    ToolParam("extension", "Uzantı filtresi", "string", required=False),
                    ToolParam("directory", "Dizin filtresi", "string", required=False),
                    ToolParam(
                        "limit", "Maks sonuç", "number", required=False, default=20
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="find_empty_files",
                description="Boş (0 byte) dosyaları bul.",
                params=[
                    ToolParam("directory", "Dizin filtresi", "string", required=False),
                ],
                danger_level="safe",
            ),
            Tool(
                name="find_empty_directories",
                description="Boş klasörleri bul (dosya ve alt klasör içermeyen).",
                params=[],
                danger_level="safe",
            ),
            Tool(
                name="get_extension_stats",
                description="Uzantı bazlı istatistikler (kaç dosya, toplam boyut).",
                params=[
                    ToolParam("directory", "Dizin filtresi", "string", required=False),
                ],
                danger_level="safe",
            ),
            Tool(
                name="find_duplicate_names",
                description="Aynı isimde birden fazla konumda olan dosyaları bul.",
                params=[
                    ToolParam("name", "Dosya adı filtresi", "string", required=False),
                    ToolParam("extension", "Uzantı filtresi", "string", required=False),
                    ToolParam(
                        "min_count",
                        "Minimum tekrar sayısı",
                        "number",
                        required=False,
                        default=2,
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="compare_files",
                description="İki dosyayı karşılaştır (aynı mı farklı mı).",
                params=[
                    ToolParam(
                        "file_path_1",
                        "Birinci dosya yolu veya adı",
                        "string",
                        required=True,
                    ),
                    ToolParam(
                        "file_path_2",
                        "İkinci dosya yolu veya adı",
                        "string",
                        required=True,
                    ),
                ],
                danger_level="safe",
            ),
            Tool(
                name="get_directory_tree",
                description="Dizin ağacını göster.",
                params=[
                    ToolParam("path", "Dizin yolu veya adı", "string", required=True),
                    ToolParam(
                        "max_depth",
                        "Kaç seviye derine in",
                        "number",
                        required=False,
                        default=3,
                    ),
                ],
                danger_level="safe",
            ),
            # --- DOSYA DÖNÜŞTÜRME ---
            Tool(
                name="convert_file",
                description="Dosyayı farklı formata dönüştür. Orijinal dosya korunur, kopya üzerinde çalışır.",
                params=[
                    ToolParam(
                        "source", "Kaynak dosya yolu veya adı", "string", required=True
                    ),
                    ToolParam(
                        "target_format",
                        "Hedef format (.pdf, .jpg, .png, .mp3, .wav, .mp4, .csv, .xlsx, .txt, .docx, .md, .html, .gif, .zip)",
                        "string",
                        required=True,
                    ),
                    ToolParam(
                        "output_dir",
                        "Çıktı dizini (varsayılan: kaynak dosyanın dizini)",
                        "string",
                        required=False,
                    ),
                    ToolParam(
                        "quality",
                        "Kalite (görsel: 1-100, ses bitrate: 128, 192, 320)",
                        "number",
                        required=False,
                    ),
                ],
                danger_level="safe",
            ),
        ]

    # ================================================================
    # DİNAMİK YOL ÇÖZÜMLEME (Tüm metotların kullandığı çekirdek)
    # ================================================================

    def _resolve_file_path(self, identifier: str) -> Optional[str]:
        """
        Verilen herhangi bir tanımlayıcıdan (isim, kısmi yol, tam yol) dosyanın
        gerçek tam yolunu bul.

        Çözümleme sırası:
            1. Tam yol ve dosya mevcutsa → doğrudan döndür
            2. Göreceli yol olarak mevcut dizinde kontrol
            3. Dosya adı olarak tam eşleşme ara (DB'de)
            4. Fuzzy arama ile bul
            5. Bulunamazsa → None

        Örnekler:
            "test.txt"           → "/home/user/Documents/test.txt"
            "/home/user/test.txt" → "/home/user/test.txt" (doğrudan)
            "config"             → "/etc/myapp/config.json" (fuzzy)
        """
        if not identifier:
            return None

        identifier = identifier.strip().strip('"').strip("'")

        # 1) Tam yol kontrolü
        p = Path(identifier)
        if p.is_absolute() and p.exists() and p.is_file():
            return str(p)

        # 2) Göreceli yol olarak mevcut dizinde kontrol
        if p.exists() and p.is_file():
            return str(p.resolve())

        # 3) DB'den tam eşleşme ile ara
        name_part = p.stem  # "test" from "test.txt"
        ext_part = p.suffix  # ".txt" from "test.txt"

        # Önce tam isim ile dene (uzantı dahil)
        if ext_part:
            exact = self.query_engine.search_files_by_exact_name(
                name=name_part, extension=ext_part
            )
            if exact:
                return exact[0]["full_path"]

        # 4) Uzantısız isim ile dene
        exact_no_ext = self.query_engine.search_files_by_exact_name(name=identifier)
        if exact_no_ext:
            return exact_no_ext[0]["full_path"]

        # 5) Fuzzy arama
        results = self.query_engine.search_files(
            name=name_part, extension=ext_part or None, limit=5
        )
        if results:
            # Tam isim eşleşmesini tercih et
            full_name = f"{name_part}{ext_part}" if ext_part else name_part
            for r in results:
                r_full = f"{r.get('name', '')}{r.get('extension', '')}"
                if r_full.lower() == full_name.lower():
                    return r["full_path"]
            # Tam eşleşme yoksa ilk sonucu döndür
            return results[0]["full_path"]

        return None

    def _resolve_dir_path(self, identifier: str) -> Optional[str]:
        """
        Verilen herhangi bir tanımlayıcıdan dizinin gerçek tam yolunu bul.

        Çözümleme sırası:
            1. Tam yol ve dizin mevcutsa → doğrudan döndür
            2. Göreceli yol kontrol
            3. Dizin adı olarak DB'de ara
            4. Bulunamazsa → None
        """
        if not identifier:
            return None

        identifier = identifier.strip().strip('"').strip("'")

        # 1) Tam yol kontrolü
        p = Path(identifier)
        if p.is_absolute() and p.exists() and p.is_dir():
            return str(p)

        # 2) Göreceli yol
        if p.exists() and p.is_dir():
            return str(p.resolve())

        # 3) DB'den ara
        dirs = self.query_engine.search_directories(name=identifier, limit=5)
        if dirs:
            # Tam isim eşleşmesi tercih et
            for d in dirs:
                if d.get("name", "").lower() == identifier.lower():
                    return d["full_path"]
            return dirs[0]["full_path"]

        return None

    def _extract_path_param(self, kwargs: dict) -> Optional[str]:
        """
        LLM'den gelen parametrelerden dosya yolunu çıkar.
        Birden fazla olası parametre adını destekler.

        Desteklenen parametre isimleri:
            path, file_path, filepath, name, file_name, filename, dosya, file
        """
        for key in (
            "path",
            "file_path",
            "filepath",
            "name",
            "file_name",
            "filename",
            "dosya",
            "file",
        ):
            if key in kwargs and kwargs[key]:
                return kwargs[key]
        return None

    def _extract_dir_param(self, kwargs: dict) -> Optional[str]:
        """
        LLM'den gelen parametrelerden dizin yolunu çıkar.

        Desteklenen parametre isimleri:
            path, dir_path, dirpath, directory, klasor, dizin, folder
        """
        for key in (
            "path",
            "dir_path",
            "dirpath",
            "directory",
            "klasor",
            "dizin",
            "folder",
        ):
            if key in kwargs and kwargs[key]:
                return kwargs[key]
        return None

    # ================================================================
    # ARAMA İŞLEMLERİ
    # ================================================================

    def search_files(
        self,
        name=None,
        extension=None,
        directory=None,
        min_size_mb=None,
        max_size_mb=None,
        limit=50,
        **kwargs,
    ) -> dict:
        """Dosya ara. Esnek parametre kabul eder."""
        if not name:
            name = kwargs.get("query") or kwargs.get("search") or kwargs.get("dosya")

        limit = max(limit, 10)

        # LLM bazen name'e "*.py" gibi wildcard gönderir — temizle
        if name:
            name = name.strip().strip("*").strip()
            if not name or name.startswith("."):
                if name and not extension:
                    extension = name if name.startswith(".") else f".{name}"
                name = None

        # extension normalize
        if extension and not extension.startswith("."):
            extension = f".{extension}"

        # Önce tam eşleşme dene
        if name:
            exact = self.query_engine.search_files_by_exact_name(
                name=name, extension=extension
            )
            if exact:
                return {"success": True, "data": exact, "count": len(exact)}

        # Tam eşleşme yoksa fuzzy ara
        results = self.query_engine.search_files(
            name=name,
            extension=extension,
            directory=directory,
            min_size_mb=min_size_mb,
            max_size_mb=max_size_mb,
            limit=limit,
        )
        return {"success": True, "data": results, "count": len(results)}

    def search_directories(
        self, name=None, parent_path=None, limit=20, **kwargs
    ) -> dict:
        """Dizin ara."""
        if not name:
            name = kwargs.get("query") or kwargs.get("search")

        results = self.query_engine.search_directories(
            name=name,
            parent_path=parent_path,
            limit=limit,
        )
        return {"success": True, "data": results, "count": len(results)}

    def find_recently_modified(
        self, days=7, extension=None, directory=None, **kwargs
    ) -> dict:
        results = self.query_engine.find_recently_modified(
            days=days,
            extension=extension,
            directory=directory,
        )
        return {"success": True, "data": results, "count": len(results)}

    def list_directory_contents(
        self, path=None, dir_path=None, extension=None, ext=None, filter=None, **kwargs
    ) -> dict:
        """Dizin içeriğini listele. Opsiyonel uzantı filtresi.

        Args:
            path/dir_path: Dizin yolu
            extension/ext/filter: Uzantı filtresi (.py, .txt, py, txt)
        """
        raw = path or dir_path or self._extract_dir_param(kwargs)
        if not raw:
            return {"success": False, "error": "Dizin belirtilmedi"}

        resolved = self._resolve_dir_path(raw)
        if not resolved:
            return {"success": False, "error": f"Dizin bulunamadı: {raw}"}

        results = self.query_engine.get_directory_contents(resolved)
        if "error" in results:
            return {"success": False, "error": results["error"]}

        # Uzantı filtresi
        ext_filter = extension or ext or filter
        if ext_filter:
            # ".py" veya "py" ikisini de kabul et
            if not ext_filter.startswith("."):
                ext_filter = "." + ext_filter
            ext_filter = ext_filter.lower()

            filtered = [
                f
                for f in results.get("files", [])
                if f.get("extension", "").lower() == ext_filter
            ]
            results["files"] = filtered
            results["filter"] = ext_filter

        return {"success": True, "data": results}

    # ================================================================
    # DOSYA OKUMA
    # ================================================================

    def read_file(
        self, path=None, file_path=None, name=None, max_lines: int = 50, **kwargs
    ) -> dict:
        """
        Dosya oku — DİNAMİK.

        Kabul edilen parametreler:
            path="test.txt"              → DB'den bulup okur
            file_path="/full/path.txt"   → doğrudan okur
            name="config"               → DB'den arayıp okur

        LLM path=, file_path=, name= ne gönderirse göndersin çalışır.
        """
        raw = path or file_path or name or self._extract_path_param(kwargs)

        if not raw:
            return {
                "success": False,
                "error": "Dosya belirtilmedi. Bir dosya adı veya yolu verin.",
            }

        # Dinamik yol çözümleme
        resolved = self._resolve_file_path(raw)
        if not resolved:
            return {
                "success": False,
                "error": f"Dosya bulunamadı: '{raw}'. "
                f"Lütfen dosyanın tam yolunu veya doğru adını verin.",
            }

        resolved_path = Path(resolved)
        if not resolved_path.is_file():
            return {"success": False, "error": f"Bu bir dosya değil: {resolved}"}

        try:
            content = resolved_path.read_text(encoding="utf-8")
            lines = content.splitlines()
            truncated = len(lines) > max_lines
            if truncated:
                lines = lines[:max_lines]
            return {
                "success": True,
                "content": "\n".join(lines),
                "total_lines": len(content.splitlines()),
                "truncated": truncated,
                "file_path": resolved,
                "message": f"{resolved_path.name} okundu ({len(content.splitlines())} satır)",
            }
        except UnicodeDecodeError:
            return {
                "success": False,
                "error": "Dosya metin formatında değil (binary olabilir)",
            }
        except Exception as e:
            return {"success": False, "error": f"Okuma hatası: {str(e)}"}

    def write_file(self, path=None, file_path=None, content="", **kwargs) -> dict:
        """
        Dosyaya yaz — DİNAMİK.
        Dosya varsa üzerine yazar, yoksa oluşturur.
        create_file'dan farkı: mevcut dosyanın üzerine yazabilir.
        """
        raw = path or file_path or self._extract_path_param(kwargs)
        content = content or kwargs.get("text", "")

        if not raw:
            return {"success": False, "error": "Dosya yolu belirtilmedi"}

        # Eğer sadece isim geldiyse ve dosya mevcutsa, önce bul
        p = Path(raw)
        if not p.is_absolute():
            # Yol içinde dizin var mı? (projeler/yedek.txt)
            if "/" in raw or "\\" in raw:
                parts = Path(raw)
                dir_part = str(parts.parent)  # "projeler"
                file_part = parts.name  # "yedek.txt"
                # Dizini çözümle
                resolved_dir = self._resolve_dir_path(dir_part)
                if resolved_dir:
                    p = Path(resolved_dir) / file_part
                else:
                    # Dizin bulunamazsa mevcut dizinde oluştur
                    resolved = self._resolve_file_path(raw)
                    if resolved:
                        p = Path(resolved)
            else:
                resolved = self._resolve_file_path(raw)
                if resolved:
                    p = Path(resolved)
                # Bulunamazsa yeni dosya olarak oluştur (mevcut dizinde)

        p.parent.mkdir(parents=True, exist_ok=True)

        try:
            p.write_text(content, encoding="utf-8")
            return {
                "success": True,
                "file_path": str(p),
                "message": f"Dosyaya yazıldı: {p.name}",
            }
        except Exception as e:
            return {"success": False, "error": f"Yazma hatası: {str(e)}"}

    def find_in_text(self, text=None, pattern=None, **kwargs) -> dict:
        """
        Metin içinde URL, IP, e-posta, domain veya özel kalıp arar.

        Desteklenen pattern türleri:
            "url"     → HTTP/HTTPS URL'leri bulur
            "ip"      → IPv4 adresleri bulur (bulamazsa domain'e fallback)
            "email"   → E-posta adresleri bulur
            "domain"  → Domain adları bulur
            "address" / "adres" → Önce IP, sonra domain, sonra URL arar (akıllı)
            "line:X"  → X metnini içeren satırları bulur
            <regex>   → Özel regex kalıbı

        Parametre esnekliği:
            text/content/data → aranacak metin
            pattern/type/find → aranacak kalıp
        """
        import re

        # Esnek parametre çözümleme
        text = (
            text or kwargs.get("content") or kwargs.get("data") or kwargs.get("input")
        )
        pattern = (
            pattern
            or kwargs.get("type")
            or kwargs.get("find")
            or kwargs.get("query")
            or kwargs.get("search")
        )

        if not text:
            return {"success": False, "error": "Aranacak metin belirtilmedi"}
        if not pattern:
            return {
                "success": False,
                "error": "Aranacak kalıp (pattern) belirtilmedi. Örn: url, ip, email veya regex",
            }

        # Eğer text aslında bir dict ise (önceki adımın çıktısı), içinden metni çıkar
        if isinstance(text, dict):
            text = (
                text.get("content")
                or text.get("data")
                or text.get("message")
                or str(text)
            )
        if isinstance(text, list):
            text = "\n".join(str(item) for item in text)

        text = str(text)

        # Önceden tanımlı regex kalıpları
        RE_IP = r"\b(?:\d{1,3}\.){3}\d{1,3}\b"
        RE_DOMAIN = r"\b(?:[a-zA-Z0-9](?:[a-zA-Z0-9-]*[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}\b"
        RE_URL = r'https?://[^\s<>"\']+'

        PATTERNS = {
            "url": RE_URL,
            "ip": RE_IP,
            "email": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
            "domain": RE_DOMAIN,
            "ipv6": r"(?:[0-9a-fA-F]{1,4}:){7}[0-9a-fA-F]{1,4}",
            "mac": r"(?:[0-9a-fA-F]{2}[:-]){5}[0-9a-fA-F]{2}",
            "phone": r"[\+]?[\d\s\-\(\)]{7,15}",
            "path": r'[A-Za-z]:\\(?:[^\\\/:*?"<>|\r\n]+\\)*[^\\\/:*?"<>|\r\n]*|/(?:[^/\0]+/)*[^/\0]+',
        }

        pattern_lower = pattern.lower().strip()

        # ── Birleşik/belirsiz pattern'ları "address"e yönlendir ──
        # LLM bazen "ip|url|domain", "ip,domain", "ip veya url" gibi gönderir
        _ADDRESS_ALIASES = {
            "address",
            "adres",
            "addr",
            "host",
            "hedef",
            "sunucu",
            "server",
            "site",
            "adress",
        }
        # Birden fazla tip içeriyorsa → address
        _MULTI_PATTERN = {"ip", "url", "domain", "host", "adres", "address"}
        pattern_tokens = set(re.split(r"[|,\s/]+", pattern_lower))
        if len(pattern_tokens & _MULTI_PATTERN) >= 2:
            pattern_lower = "address"

        if pattern_lower in _ADDRESS_ALIASES:
            return self._find_address_smart(text, RE_IP, RE_DOMAIN, RE_URL)

        # ── Satır bazlı arama: "line:kelime" ──
        if pattern_lower.startswith("line:"):
            search_term = pattern[5:].strip()
            matches = []
            for i, line in enumerate(text.splitlines(), 1):
                if search_term.lower() in line.lower():
                    matches.append({"line": i, "text": line.strip()})
            return {
                "success": True,
                "data": matches,
                "count": len(matches),
                "pattern_type": "line_search",
                "message": f"'{search_term}' için {len(matches)} satır eşleşti",
            }

        # ── Bilinen pattern tipi mi? ──
        if pattern_lower in PATTERNS:
            regex = PATTERNS[pattern_lower]
            pattern_type = pattern_lower
        else:
            # Özel regex olarak kullan
            try:
                re.compile(pattern)
                regex = pattern
                pattern_type = "custom_regex"
            except re.error:
                # Regex değilse düz metin araması yap
                matches = []
                for i, line in enumerate(text.splitlines(), 1):
                    if pattern.lower() in line.lower():
                        matches.append({"line": i, "text": line.strip()})
                return {
                    "success": True,
                    "data": [line["text"] for line in matches] if matches else [],
                    "count": len(matches),
                    "pattern_type": "text_search",
                    "message": f"'{pattern}' için {len(matches)} eşleşme bulundu",
                }

        # ── Regex ile bul ──
        found = re.findall(regex, text)
        unique = list(dict.fromkeys(found))

        # ── IP bulunamadıysa domain'e fallback ──
        if not unique and pattern_lower == "ip":
            domain_found = re.findall(RE_DOMAIN, text)
            domain_unique = list(dict.fromkeys(domain_found))
            if domain_unique:
                return {
                    "success": True,
                    "data": domain_unique,
                    "count": len(domain_unique),
                    "pattern_type": "domain (ip bulunamadı, domain'e geçildi)",
                    "message": f"IP bulunamadı, {len(domain_unique)} domain bulundu",
                }

        # ── URL bulunamadıysa domain'e fallback ──
        if not unique and pattern_lower == "url":
            domain_found = re.findall(RE_DOMAIN, text)
            domain_unique = list(dict.fromkeys(domain_found))
            if domain_unique:
                return {
                    "success": True,
                    "data": domain_unique,
                    "count": len(domain_unique),
                    "pattern_type": "domain (url bulunamadı, domain'e geçildi)",
                    "message": f"URL bulunamadı, {len(domain_unique)} domain bulundu",
                }

        return {
            "success": True,
            "data": unique,
            "count": len(unique),
            "total_with_duplicates": len(found),
            "pattern_type": pattern_type,
            "message": f"{len(unique)} benzersiz {pattern_type} bulundu"
            if unique
            else f"Hiç {pattern_type} bulunamadı",
        }

    def _find_address_smart(
        self, text: str, re_ip: str, re_domain: str, re_url: str
    ) -> dict:
        """
        Akıllı adres arama: metin içinde ne tür adres varsa onu bul.
        Öncelik sırası: IP → Domain → URL → satır bazlı fallback
        """
        import re

        all_found = []
        found_types = []

        # 1) IP adresleri
        ips = list(dict.fromkeys(re.findall(re_ip, text)))
        if ips:
            all_found.extend(ips)
            found_types.append(f"{len(ips)} IP")

        # 2) Domain adresleri (IP olmayanlar)
        domains = list(dict.fromkeys(re.findall(re_domain, text)))
        # IP'leri domain listesinden çıkar (çakışmayı önle)
        domains = [d for d in domains if d not in ips]
        if domains:
            all_found.extend(domains)
            found_types.append(f"{len(domains)} domain")

        # 3) URL'ler (zaten domain'e dahil olanları tekrarlama)
        urls = list(dict.fromkeys(re.findall(re_url, text)))
        new_urls = [
            u for u in urls if not any(u.endswith(d) or d in u for d in domains + ips)
        ]
        if new_urls:
            all_found.extend(new_urls)
            found_types.append(f"{len(new_urls)} URL")

        # 4) Hiçbir şey bulunamadıysa → boş dön (eski: tüm satırları döndürüyordu)
        type_summary = ", ".join(found_types) if found_types else "hiçbir şey"

        return {
            "success": True,
            "data": all_found,
            "count": len(all_found),
            "pattern_type": "address_smart",
            "message": f"Bulunan: {type_summary}"
            if all_found
            else "Hiç adres bulunamadı",
        }

    # ================================================================
    # KOPYALAMA / TAŞIMA
    # ================================================================

    def copy_file(self, source=None, destination=None, **kwargs) -> dict:
        """Dosya kopyala — kaynak dinamik çözümlenir."""
        source = source or kwargs.get("src") or kwargs.get("from")
        destination = (
            destination
            or kwargs.get("dest")
            or kwargs.get("to")
            or kwargs.get("target")
        )

        if not source:
            return {"success": False, "error": "Kaynak dosya belirtilmedi"}
        if not destination:
            return {"success": False, "error": "Hedef belirtilmedi"}

        resolved_src = self._resolve_file_path(source)
        if not resolved_src:
            return {"success": False, "error": f"Kaynak dosya bulunamadı: {source}"}

        src = Path(resolved_src)
        dst = Path(destination)

        if not dst.is_absolute():
            resolved_dst_dir = self._resolve_dir_path(destination)
            if resolved_dst_dir:
                dst = Path(resolved_dst_dir) / src.name
            else:
                dst = src.parent / destination

        if dst.is_dir():
            dst = dst / src.name

        dst.parent.mkdir(parents=True, exist_ok=True)

        try:
            shutil.copy2(str(src), str(dst))
            return {
                "success": True,
                "source": str(src),
                "destination": str(dst),
                "message": f"Dosya kopyalandı: {src.name}",
            }
        except Exception as e:
            return {"success": False, "error": f"Kopyalama hatası: {str(e)}"}

    def move_file(self, source=None, destination=None, **kwargs) -> dict:
        """Dosya taşı — kaynak dinamik çözümlenir."""
        source = source or kwargs.get("src") or kwargs.get("from")
        destination = (
            destination
            or kwargs.get("dest")
            or kwargs.get("to")
            or kwargs.get("target")
        )

        if not source:
            return {"success": False, "error": "Kaynak dosya belirtilmedi"}
        if not destination:
            return {"success": False, "error": "Hedef belirtilmedi"}

        resolved_src = self._resolve_file_path(source)
        if not resolved_src:
            return {"success": False, "error": f"Kaynak bulunamadı: {source}"}

        src = Path(resolved_src)
        dst = Path(destination)

        if not dst.is_absolute():
            resolved_dst_dir = self._resolve_dir_path(destination)
            if resolved_dst_dir:
                dst = Path(resolved_dst_dir) / src.name
            else:
                dst = src.parent / destination

        if dst.is_dir():
            dst = dst / src.name

        dst.parent.mkdir(parents=True, exist_ok=True)

        try:
            shutil.move(str(src), str(dst))
            return {
                "success": True,
                "source": str(src),
                "destination": str(dst),
                "message": f"Dosya taşındı: {src.name}",
            }
        except Exception as e:
            return {"success": False, "error": f"Taşıma hatası: {str(e)}"}

    # ================================================================
    # OLUŞTURMA
    # ================================================================

    def create_file(self, path=None, file_path=None, content="", **kwargs) -> dict:
        """Yeni dosya oluştur — esnek parametre."""
        raw = path or file_path or self._extract_path_param(kwargs)
        if not raw:
            return {"success": False, "error": "Dosya yolu belirtilmedi"}

        content = content or kwargs.get("text", "")
        p = Path(raw)

        # Göreli yolda dizin varsa çözümle (projeler/yeni.txt)
        if not p.is_absolute() and ("/" in raw or "\\" in raw):
            dir_part = str(Path(raw).parent)
            file_part = Path(raw).name
            resolved_dir = self._resolve_dir_path(dir_part)
            if resolved_dir:
                p = Path(resolved_dir) / file_part

        if p.exists():
            return {"success": False, "error": f"Dosya zaten var: {raw}"}

        p.parent.mkdir(parents=True, exist_ok=True)

        try:
            p.write_text(content, encoding="utf-8")
            return {
                "success": True,
                "file_path": str(p),
                "message": f"Dosya oluşturuldu: {p.name}",
            }
        except Exception as e:
            return {"success": False, "error": f"Oluşturma hatası: {str(e)}"}

    def create_directory(self, path=None, dir_path=None, **kwargs) -> dict:
        """Yeni dizin oluştur — esnek parametre."""
        raw = path or dir_path or self._extract_dir_param(kwargs)
        if not raw:
            return {"success": False, "error": "Dizin yolu belirtilmedi"}

        p = Path(raw)

        if p.exists():
            return {"success": False, "error": f"Dizin zaten var: {raw}"}

        try:
            p.mkdir(parents=True, exist_ok=True)
            return {
                "success": True,
                "dir_path": str(p),
                "message": f"Dizin oluşturuldu: {p.name}",
            }
        except Exception as e:
            return {"success": False, "error": f"Oluşturma hatası: {str(e)}"}

    # ================================================================
    # SİLME
    # ================================================================

    def delete_file(self, path=None, file_path=None, **kwargs) -> dict:
        """Dosya sil — dinamik çözümleme."""
        raw = path or file_path or self._extract_path_param(kwargs)
        if not raw:
            return {"success": False, "error": "Silinecek dosya belirtilmedi"}

        resolved = self._resolve_file_path(raw)
        if not resolved:
            return {"success": False, "error": f"Dosya bulunamadı: {raw}"}

        try:
            Path(resolved).unlink()
            return {
                "success": True,
                "file_path": resolved,
                "message": f"Dosya silindi: {Path(resolved).name}",
            }
        except Exception as e:
            return {"success": False, "error": f"Silme hatası: {str(e)}"}

    def delete_directory(self, path=None, dir_path=None, **kwargs) -> dict:
        """Dizin sil — dinamik çözümleme."""
        raw = path or dir_path or self._extract_dir_param(kwargs)
        if not raw:
            return {"success": False, "error": "Silinecek dizin belirtilmedi"}

        resolved = self._resolve_dir_path(raw)
        if not resolved:
            return {"success": False, "error": f"Dizin bulunamadı: {raw}"}

        try:
            shutil.rmtree(resolved)
            return {
                "success": True,
                "dir_path": resolved,
                "message": f"Dizin silindi: {Path(resolved).name}",
            }
        except Exception as e:
            return {"success": False, "error": f"Silme hatası: {str(e)}"}

    # ================================================================
    # YENİDEN ADLANDIRMA
    # ================================================================

    def rename(self, path=None, new_name=None, **kwargs) -> dict:
        """Dosya/dizin yeniden adlandır — dinamik çözümleme."""
        raw = path or self._extract_path_param(kwargs)
        new_name = new_name or kwargs.get("new") or kwargs.get("to")

        if not raw:
            return {"success": False, "error": "Dosya/dizin belirtilmedi"}
        if not new_name:
            return {"success": False, "error": "Yeni isim belirtilmedi"}

        # Önce dosya olarak dene
        resolved = self._resolve_file_path(raw)
        # Dosya bulunamazsa dizin olarak dene
        if not resolved:
            resolved = self._resolve_dir_path(raw)
        if not resolved:
            return {"success": False, "error": f"Bulunamadı: {raw}"}

        src = Path(resolved)
        dst = src.parent / new_name

        if dst.exists():
            return {"success": False, "error": f"Bu isimde zaten var: {new_name}"}

        try:
            src.rename(dst)
            return {
                "success": True,
                "old_path": str(src),
                "new_path": str(dst),
                "message": f"Yeniden adlandırıldı: {src.name} -> {new_name}",
            }
        except Exception as e:
            return {"success": False, "error": f"Yeniden adlandırma hatası: {str(e)}"}

    # ================================================================
    # BİLGİ
    # ================================================================

    def get_file_info(self, path=None, file_path=None, **kwargs) -> dict:
        """Dosya bilgisi — dinamik çözümleme."""
        raw = path or file_path or self._extract_path_param(kwargs)
        if not raw:
            return {"success": False, "error": "Dosya belirtilmedi"}

        resolved = self._resolve_file_path(raw)
        if not resolved:
            return {"success": False, "error": f"Bulunamadı: {raw}"}

        p = Path(resolved)
        stat = p.stat()
        size = stat.st_size

        return {
            "success": True,
            "data": {
                "name": p.name,
                "extension": p.suffix,
                "full_path": str(p),
                "directory": str(p.parent),
                "size_bytes": size,
                "size_human": self._human_size(size),
                "created": datetime.fromtimestamp(stat.st_ctime).strftime(
                    "%Y-%m-%d %H:%M"
                ),
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime(
                    "%Y-%m-%d %H:%M"
                ),
                "is_file": p.is_file(),
                "is_dir": p.is_dir(),
            },
        }

    def find_directory_size(self, path=None, dir_path=None, **kwargs) -> dict:
        """Dizin boyutu — dinamik çözümleme."""
        raw = path or dir_path or self._extract_dir_param(kwargs)
        if not raw:
            return {"success": False, "error": "Dizin belirtilmedi"}

        resolved = self._resolve_dir_path(raw)
        if not resolved:
            return {"success": False, "error": f"Dizin bulunamadı: {raw}"}

        result = self.query_engine.get_directory_size(resolved)
        return {"success": True, "data": result}

    # ================================================================
    # İÇERİK İŞLEMLERİ
    # ================================================================

    def search_in_files(
        self,
        query=None,
        directory=None,
        extension=None,
        is_regex=False,
        limit=20,
        **kwargs,
    ) -> dict:
        """Dosyaların içinde metin ara."""
        import re

        query = (
            query or kwargs.get("text") or kwargs.get("search") or kwargs.get("pattern")
        )
        if not query:
            return {"success": False, "error": "Aranacak metin belirtilmedi"}

        results = []
        search_files = self.query_engine.search_files(
            extension=extension, directory=directory, limit=500
        )
        pattern = re.compile(query) if is_regex else None

        for f in search_files:
            try:
                fpath = Path(f["full_path"])
                content = fpath.read_text(encoding="utf-8", errors="ignore")
                matches = []
                for i, line in enumerate(content.splitlines(), 1):
                    if pattern:
                        if pattern.search(line):
                            matches.append({"line": i, "text": line.strip()})
                    else:
                        if query.lower() in line.lower():
                            matches.append({"line": i, "text": line.strip()})
                if matches:
                    results.append(
                        {
                            "file": f["full_path"],
                            "name": f"{f['name']}{f['extension']}",
                            "matches": matches[:5],
                            "total_matches": len(matches),
                        }
                    )
                    if len(results) >= limit:
                        break
            except Exception:
                continue

        return {"success": True, "data": results, "count": len(results)}

    def append_to_file(self, path=None, file_path=None, content=None, **kwargs) -> dict:
        """Dosyaya içerik ekle — dinamik çözümleme. Dosya yoksa oluşturur."""
        raw = path or file_path or self._extract_path_param(kwargs)
        content = content or kwargs.get("text", "")

        if not raw:
            return {"success": False, "error": "Dosya belirtilmedi"}
        if not content:
            return {"success": False, "error": "Eklenecek içerik belirtilmedi"}

        # Önce mevcut dosyayı bul
        resolved = self._resolve_file_path(raw)

        if not resolved:
            # Dosya yok → write_file gibi oluştur
            p = Path(raw)
            if not p.is_absolute():
                # Dizin yolu varsa çözümle (projeler/birlesik.txt)
                if "/" in raw or "\\" in raw:
                    dir_part = str(Path(raw).parent)
                    file_part = Path(raw).name
                    resolved_dir = self._resolve_dir_path(dir_part)
                    if resolved_dir:
                        p = Path(resolved_dir) / file_part
                # Dizin yolu yoksa mevcut dizinde oluştur
            p.parent.mkdir(parents=True, exist_ok=True)
            resolved = str(p)

        try:
            with open(resolved, "a", encoding="utf-8") as f:
                f.write(content + "\n")
            return {
                "success": True,
                "message": f"İçerik eklendi: {Path(resolved).name}",
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def replace_text(
        self,
        path=None,
        file_path=None,
        old_text=None,
        new_text=None,
        is_regex=False,
        **kwargs,
    ) -> dict:
        """Dosyada metin bul-değiştir — dinamik çözümleme."""
        import re

        raw = path or file_path or self._extract_path_param(kwargs)
        old_text = old_text or kwargs.get("find") or kwargs.get("search")
        new_text = new_text if new_text is not None else kwargs.get("replace", "")

        if not raw:
            return {"success": False, "error": "Dosya belirtilmedi"}
        if not old_text:
            return {"success": False, "error": "Değiştirilecek metin belirtilmedi"}

        resolved = self._resolve_file_path(raw)
        if not resolved:
            return {"success": False, "error": f"Dosya bulunamadı: {raw}"}

        try:
            p = Path(resolved)
            content = p.read_text(encoding="utf-8")
            if is_regex:
                new_content = re.sub(old_text, new_text, content)
                count = len(re.findall(old_text, content))
            else:
                count = content.count(old_text)
                new_content = content.replace(old_text, new_text)
            if count == 0:
                return {"success": False, "error": f"'{old_text}' bulunamadı"}
            p.write_text(new_content, encoding="utf-8")
            return {"success": True, "message": f"{count} değişiklik yapıldı: {p.name}"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def count_lines_words(self, path=None, file_path=None, **kwargs) -> dict:
        """Satır/kelime/karakter say — dinamik çözümleme."""
        raw = path or file_path or self._extract_path_param(kwargs)
        if not raw:
            return {"success": False, "error": "Dosya belirtilmedi"}

        resolved = self._resolve_file_path(raw)
        if not resolved:
            return {"success": False, "error": f"Dosya bulunamadı: {raw}"}

        try:
            p = Path(resolved)
            content = p.read_text(encoding="utf-8")
            lines = content.splitlines()
            words = content.split()
            return {
                "success": True,
                "data": {
                    "file": p.name,
                    "lines": len(lines),
                    "words": len(words),
                    "characters": len(content),
                },
                "message": f"{p.name}: {len(lines)} satır, {len(words)} kelime, {len(content)} karakter",
            }
        except UnicodeDecodeError:
            return {"success": False, "error": "Dosya metin formatında değil"}

    # ================================================================
    # BÜYÜK / KÜÇÜK / BOŞ DOSYALAR
    # ================================================================

    def find_largest_files(
        self, min_size_mb=None, extension=None, directory=None, limit=10, **kwargs
    ) -> dict:
        results = self.query_engine.find_largest_files(
            extension=extension, directory=directory, limit=limit
        )
        if min_size_mb is not None:
            min_bytes = int(min_size_mb * 1024 * 1024)
            results = [r for r in results if r.get("size_bytes", 0) >= min_bytes]
        return {"success": True, "data": results, "count": len(results)}

    def find_small_files(
        self, max_size_mb=None, extension=None, directory=None, limit=20, **kwargs
    ) -> dict:
        results = self.query_engine.find_small_files(
            extension=extension, directory=directory, limit=limit
        )
        if max_size_mb is not None:
            max_bytes = int(max_size_mb * 1024 * 1024)
            results = [r for r in results if r.get("size_bytes", 0) <= max_bytes]
        return {"success": True, "data": results, "count": len(results)}

    def find_empty_files(self, directory=None, **kwargs) -> dict:
        results = self.query_engine.find_empty_files(directory=directory)
        return {"success": True, "data": results, "count": len(results)}

    def find_empty_directories(self, **kwargs) -> dict:
        results = self.query_engine.find_empty_directories()
        return {"success": True, "data": results, "count": len(results)}

    def get_extension_stats(self, directory=None, **kwargs) -> dict:
        results = self.query_engine.get_extension_stats(directory=directory)
        return {"success": True, "data": results, "count": len(results)}

    def find_duplicate_names(
        self, extension=None, min_count=2, name=None, **kwargs
    ) -> dict:
        results = self.query_engine.find_duplicate_names(
            extension=extension, min_count=min_count
        )
        if name:
            results = [r for r in results if name.lower() in r["name"].lower()]
        return {"success": True, "data": results, "count": len(results)}

    def compare_files(self, file_path_1=None, file_path_2=None, **kwargs) -> dict:
        """İki dosyayı karşılaştır — dinamik çözümleme."""
        import hashlib

        fp1 = file_path_1 or kwargs.get("path1") or kwargs.get("first")
        fp2 = file_path_2 or kwargs.get("path2") or kwargs.get("second")

        if not fp1 or not fp2:
            return {"success": False, "error": "İki dosya belirtilmeli"}

        resolved1 = self._resolve_file_path(fp1)
        resolved2 = self._resolve_file_path(fp2)

        if not resolved1:
            return {"success": False, "error": f"Bulunamadı: {fp1}"}
        if not resolved2:
            return {"success": False, "error": f"Bulunamadı: {fp2}"}

        p1, p2 = Path(resolved1), Path(resolved2)
        s1, s2 = p1.stat().st_size, p2.stat().st_size

        if s1 != s2:
            return {
                "success": True,
                "data": {"identical": False, "reason": f"Boyut farklı: {s1} vs {s2}"},
            }

        h1 = hashlib.md5(p1.read_bytes()).hexdigest()
        h2 = hashlib.md5(p2.read_bytes()).hexdigest()
        identical = h1 == h2

        return {
            "success": True,
            "data": {
                "identical": identical,
                "reason": "Aynı dosya" if identical else "İçerik farklı",
                "hash_1": h1,
                "hash_2": h2,
            },
            "message": "Dosyalar aynı" if identical else "Dosyalar farklı",
        }

    def get_directory_tree(
        self, path=None, dir_path=None, max_depth=3, **kwargs
    ) -> dict:
        """Dizin ağacı — dinamik çözümleme."""
        raw = path or dir_path or self._extract_dir_param(kwargs)
        if not raw:
            return {"success": False, "error": "Dizin belirtilmedi"}

        resolved = self._resolve_dir_path(raw)
        if not resolved:
            return {"success": False, "error": f"Dizin bulunamadı: {raw}"}

        result = self.query_engine.get_directory_tree(resolved, max_depth=max_depth)
        if "error" in result:
            return {"success": False, "error": result["error"]}
        return {"success": True, "data": result}

    # ================================================================
    # DOSYA DÖNÜŞTÜRME
    # ================================================================

    # Desteklenen dönüşüm çiftleri → handler metot adı
    _CONVERTERS = {
        # Doküman
        (".txt", ".pdf"): "_conv_text_to_pdf",
        (".md", ".pdf"): "_conv_text_to_pdf",
        (".html", ".pdf"): "_conv_html_to_pdf",
        (".csv", ".xlsx"): "_conv_csv_to_xlsx",
        (".xlsx", ".csv"): "_conv_xlsx_to_csv",
        (".txt", ".html"): "_conv_text_to_html",
        (".md", ".html"): "_conv_md_to_html",
        (".txt", ".md"): "_conv_rename_ext",
        (".md", ".txt"): "_conv_rename_ext",
        (".txt", ".docx"): "_conv_text_to_docx",
        (".docx", ".txt"): "_conv_docx_to_txt",
        (".docx", ".pdf"): "_conv_docx_to_pdf",
        # Görsel
        (".png", ".jpg"): "_conv_image",
        (".jpg", ".png"): "_conv_image",
        (".jpeg", ".png"): "_conv_image",
        (".png", ".jpeg"): "_conv_image",
        (".webp", ".png"): "_conv_image",
        (".webp", ".jpg"): "_conv_image",
        (".bmp", ".png"): "_conv_image",
        (".bmp", ".jpg"): "_conv_image",
        (".png", ".gif"): "_conv_image",
        (".jpg", ".gif"): "_conv_image",
        (".png", ".webp"): "_conv_image",
        (".jpg", ".webp"): "_conv_image",
        (".png", ".bmp"): "_conv_image",
        # Ses
        (".wav", ".mp3"): "_conv_audio",
        (".mp3", ".wav"): "_conv_audio",
        (".ogg", ".mp3"): "_conv_audio",
        (".flac", ".mp3"): "_conv_audio",
        (".m4a", ".mp3"): "_conv_audio",
        (".mp3", ".ogg"): "_conv_audio",
        (".wav", ".ogg"): "_conv_audio",
        (".flac", ".wav"): "_conv_audio",
        # Video
        (".mp4", ".mp3"): "_conv_video_to_audio",
        (".mp4", ".wav"): "_conv_video_to_audio",
        (".mkv", ".mp4"): "_conv_video",
        (".avi", ".mp4"): "_conv_video",
        (".webm", ".mp4"): "_conv_video",
        (".mov", ".mp4"): "_conv_video",
        (".mp4", ".gif"): "_conv_video_to_gif",
        # Sıkıştırma
        (".zip", ".dir"): "_conv_unzip",  # özel: zip çıkar
    }

    def convert_file(
        self,
        source=None,
        target_format=None,
        output_dir=None,
        quality=None,
        path=None,
        **kwargs,
    ) -> dict:
        """Dosyayı farklı formata dönüştür. Orijinal korunur, kopya üzerinde çalışır."""
        source = source or path or kwargs.get("file_path") or kwargs.get("file")
        target_format = target_format or kwargs.get("format") or kwargs.get("to")

        if not source:
            return {"success": False, "error": "Kaynak dosya belirtilmedi"}
        if not target_format:
            return {"success": False, "error": "Hedef format belirtilmedi"}

        # Quality normalize — LLM bazen "high", "low" gibi string verir
        if isinstance(quality, str):
            _quality_map = {"high": 95, "medium": 75, "low": 50}
            try:
                quality = int(quality)
            except ValueError:
                quality = _quality_map.get(quality.lower())

        # Hedef formatı normalize et
        target_format = target_format.strip().lower()
        if not target_format.startswith("."):
            target_format = f".{target_format}"

        # Kaynak dosyayı bul
        resolved = self._resolve_file_path(source)

        # _resolve_file_path bulamadıysa ek aramalar yap
        if not resolved:
            src_p = Path(source)

            # 1) Göreceli yol olarak izlenen dizinlerde ara
            #    "transform_test/test_resim.png" → Desktop/projeler/transform_test/test_resim.png
            if "/" in source or "\\" in source:
                # Parçalara ayır — son kısmı dosya adı, öncesi dizin
                dir_part = str(src_p.parent)  # "transform_test"
                file_name = src_p.name  # "test_resim.png"

                resolved_dir = self._resolve_dir_path(dir_part)
                if resolved_dir:
                    candidate = Path(resolved_dir) / file_name
                    if candidate.exists():
                        resolved = str(candidate)

            # 2) Sadece dosya adıyla search_files'ta ara
            if not resolved and src_p.suffix:
                results = self.query_engine.search_files(
                    name=src_p.stem, extension=src_p.suffix, limit=5
                )
                if results:
                    # Tam isim eşleşmesini tercih et
                    target_name = src_p.name.lower()
                    for r in results:
                        r_name = f"{r.get('name', '')}{r.get('extension', '')}".lower()
                        if r_name == target_name:
                            resolved = r["full_path"]
                            break
                    # Tam eşleşme yoksa ilk sonucu al
                    if not resolved:
                        resolved = results[0]["full_path"]

        if not resolved:
            return {"success": False, "error": f"Kaynak dosya bulunamadı: {source}"}

        src_path = Path(resolved)
        src_ext = src_path.suffix.lower()

        # Aynı format kontrolü
        if src_ext == target_format:
            return {
                "success": False,
                "error": f"Dosya zaten {target_format} formatında",
            }

        # Handler bul
        handler_name = self._CONVERTERS.get((src_ext, target_format))
        if not handler_name:
            supported = self._get_supported_conversions(src_ext)
            return {
                "success": False,
                "error": f"{src_ext} → {target_format} dönüşümü desteklenmiyor",
                "supported": supported,
            }

        # Çıktı dizinini belirle
        if output_dir:
            out_dir = self._resolve_dir_path(output_dir)
            if not out_dir:
                out_dir = output_dir
            Path(out_dir).mkdir(parents=True, exist_ok=True)
        else:
            out_dir = str(src_path.parent)

        # Hedef dosya adı
        target_name = src_path.stem + target_format
        target_path = Path(out_dir) / target_name

        # Hedef varsa numaralı ad ver
        counter = 1
        while target_path.exists():
            target_name = f"{src_path.stem}_{counter}{target_format}"
            target_path = Path(out_dir) / target_name
            counter += 1

        # Dosya boyut kontrolü (100MB üzeri uyarı)
        src_size = src_path.stat().st_size
        if src_size > 100 * 1024 * 1024:
            size_mb = src_size / (1024 * 1024)
            return {
                "success": False,
                "error": f"Dosya çok büyük ({size_mb:.1f} MB). Güvenlik sınırı: 100MB",
            }

        # Handler'ı çağır
        handler = getattr(self, handler_name, None)
        if not handler:
            return {"success": False, "error": f"Handler bulunamadı: {handler_name}"}

        try:
            result = handler(src_path, target_path, quality=quality)
            if result.get("success"):
                result["source"] = str(src_path)
                result["target"] = str(target_path)
                result["source_format"] = src_ext
                result["target_format"] = target_format
                result["source_size"] = src_size
                if target_path.exists():
                    result["target_size"] = target_path.stat().st_size
            return result
        except Exception as e:
            # Hata durumunda oluşan dosyayı temizle
            if target_path.exists():
                target_path.unlink()
            return {"success": False, "error": f"Dönüştürme hatası: {str(e)}"}

    def _get_supported_conversions(self, src_ext: str) -> list:
        """Bir kaynak format için desteklenen hedef formatları döndür."""
        return [t for (s, t) in self._CONVERTERS.keys() if s == src_ext]

    # ── HANDLER: Sadece uzantı değiştir ──
    @staticmethod
    def _conv_rename_ext(src: Path, target: Path, **kw) -> dict:
        shutil.copy2(str(src), str(target))
        return {"success": True, "message": f"Dönüştürüldü: {target.name}"}

    # ── HANDLER: Text/Markdown → PDF ──
    @staticmethod
    def _conv_text_to_pdf(src: Path, target: Path, **kw) -> dict:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import cm
        except ImportError:
            return {
                "success": False,
                "error": "reportlab kütüphanesi gerekli: pip install reportlab",
            }

        text = src.read_text(encoding="utf-8", errors="replace")
        c = canvas.Canvas(str(target), pagesize=A4)
        width, height = A4
        margin = 2 * cm
        y = height - margin
        line_height = 14

        # Font ayarla (Türkçe karakter desteği)
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os

            # Windows'ta Segoe UI veya Arial dene
            for font_path in [
                "C:/Windows/Fonts/arial.ttf",
                "C:/Windows/Fonts/segoeui.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            ]:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont("CustomFont", font_path))
                    c.setFont("CustomFont", 10)
                    break
            else:
                c.setFont("Helvetica", 10)
        except Exception:
            c.setFont("Helvetica", 10)

        for line in text.split("\n"):
            if y < margin:
                c.showPage()
                y = height - margin
                try:
                    c.setFont("CustomFont", 10)
                except Exception:
                    c.setFont("Helvetica", 10)
            c.drawString(margin, y, line[:120])  # Satır uzunluk sınırı
            y -= line_height

        c.save()
        return {"success": True, "message": f"PDF oluşturuldu: {target.name}"}

    # ── HANDLER: HTML → PDF ──
    @staticmethod
    def _conv_html_to_pdf(src: Path, target: Path, **kw) -> dict:
        # weasyprint varsa kullan, yoksa basit reportlab
        try:
            from weasyprint import HTML

            HTML(filename=str(src)).write_pdf(str(target))
            return {"success": True, "message": f"PDF oluşturuldu: {target.name}"}
        except ImportError:
            pass
        # Fallback: HTML'i text olarak oku, PDF yap
        return FileOperationsModule._conv_text_to_pdf(src, target, **kw)

    # ── HANDLER: Text → HTML ──
    @staticmethod
    def _conv_text_to_html(src: Path, target: Path, **kw) -> dict:
        text = src.read_text(encoding="utf-8", errors="replace")
        lines = text.split("\n")
        html_lines = [f"<p>{line}</p>" if line.strip() else "<br>" for line in lines]
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{src.stem}</title></head>
<body>{"".join(html_lines)}</body></html>"""
        target.write_text(html, encoding="utf-8")
        return {"success": True, "message": f"HTML oluşturuldu: {target.name}"}

    # ── HANDLER: Markdown → HTML ──
    @staticmethod
    def _conv_md_to_html(src: Path, target: Path, **kw) -> dict:
        text = src.read_text(encoding="utf-8", errors="replace")
        try:
            import markdown

            html_body = markdown.markdown(text)
        except ImportError:
            # Basit fallback
            html_body = text.replace("\n", "<br>")
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{src.stem}</title></head>
<body>{html_body}</body></html>"""
        target.write_text(html, encoding="utf-8")
        return {"success": True, "message": f"HTML oluşturuldu: {target.name}"}

    # ── HANDLER: Text → DOCX ──
    @staticmethod
    def _conv_text_to_docx(src: Path, target: Path, **kw) -> dict:
        try:
            from docx import Document
        except ImportError:
            return {
                "success": False,
                "error": "python-docx kütüphanesi gerekli: pip install python-docx",
            }
        text = src.read_text(encoding="utf-8", errors="replace")
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(target))
        return {"success": True, "message": f"DOCX oluşturuldu: {target.name}"}

    # ── HANDLER: DOCX → Text ──
    @staticmethod
    def _conv_docx_to_txt(src: Path, target: Path, **kw) -> dict:
        try:
            from docx import Document
        except ImportError:
            return {
                "success": False,
                "error": "python-docx kütüphanesi gerekli: pip install python-docx",
            }
        doc = Document(str(src))
        lines = [p.text for p in doc.paragraphs]
        target.write_text("\n".join(lines), encoding="utf-8")
        return {"success": True, "message": f"TXT oluşturuldu: {target.name}"}

    # ── HANDLER: DOCX → PDF ──
    @staticmethod
    def _conv_docx_to_pdf(src: Path, target: Path, **kw) -> dict:
        try:
            from docx import Document
        except ImportError:
            return {"success": False, "error": "python-docx kütüphanesi gerekli"}
        # DOCX → geçici TXT → PDF
        doc = Document(str(src))
        lines = [p.text for p in doc.paragraphs]
        import tempfile

        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        ) as tmp:
            tmp.write("\n".join(lines))
            tmp_path = tmp.name
        try:
            result = FileOperationsModule._conv_text_to_pdf(
                Path(tmp_path), target, **kw
            )
        finally:
            Path(tmp_path).unlink(missing_ok=True)
        return result

    # ── HANDLER: CSV → XLSX ──
    @staticmethod
    def _conv_csv_to_xlsx(src: Path, target: Path, **kw) -> dict:
        try:
            from openpyxl import Workbook
        except ImportError:
            return {
                "success": False,
                "error": "openpyxl kütüphanesi gerekli: pip install openpyxl",
            }
        import csv

        wb = Workbook()
        ws = wb.active
        text = src.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(text.splitlines())
        for row in reader:
            ws.append(row)
        wb.save(str(target))
        return {"success": True, "message": f"XLSX oluşturuldu: {target.name}"}

    # ── HANDLER: XLSX → CSV ──
    @staticmethod
    def _conv_xlsx_to_csv(src: Path, target: Path, **kw) -> dict:
        try:
            from openpyxl import load_workbook
        except ImportError:
            return {"success": False, "error": "openpyxl kütüphanesi gerekli"}
        import csv

        wb = load_workbook(str(src), read_only=True)
        ws = wb.active
        with open(str(target), "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(row)
        wb.close()
        return {"success": True, "message": f"CSV oluşturuldu: {target.name}"}

    # ── HANDLER: Görsel dönüştürme ──
    @staticmethod
    def _conv_image(src: Path, target: Path, quality=None, **kw) -> dict:
        try:
            from PIL import Image
        except ImportError:
            return {
                "success": False,
                "error": "Pillow kütüphanesi gerekli: pip install Pillow",
            }
        img = Image.open(str(src))
        # RGBA → RGB (JPEG için gerekli)
        if target.suffix.lower() in (".jpg", ".jpeg") and img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        save_kw = {}
        if quality and target.suffix.lower() in (".jpg", ".jpeg", ".webp"):
            save_kw["quality"] = int(quality)
        img.save(str(target), **save_kw)
        return {"success": True, "message": f"Görsel dönüştürüldü: {target.name}"}

    # ── HANDLER: Ses dönüştürme (ffmpeg) ──
    @staticmethod
    def _conv_audio(src: Path, target: Path, quality=None, **kw) -> dict:
        ffmpeg = FileOperationsModule._find_ffmpeg()
        if not ffmpeg:
            return {
                "success": False,
                "error": "ffmpeg bulunamadı. Lütfen ffmpeg kurun.",
            }
        import subprocess

        cmd = [ffmpeg, "-i", str(src), "-y"]
        if quality and target.suffix.lower() == ".mp3":
            cmd += ["-b:a", f"{int(quality)}k"]
        cmd.append(str(target))
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            return {"success": False, "error": f"ffmpeg hatası: {result.stderr[:300]}"}
        return {"success": True, "message": f"Ses dönüştürüldü: {target.name}"}

    # ── HANDLER: Video → Ses ──
    @staticmethod
    def _conv_video_to_audio(src: Path, target: Path, quality=None, **kw) -> dict:
        ffmpeg = FileOperationsModule._find_ffmpeg()
        if not ffmpeg:
            return {"success": False, "error": "ffmpeg bulunamadı."}
        import subprocess

        cmd = [ffmpeg, "-i", str(src), "-vn", "-y"]
        if quality and target.suffix.lower() == ".mp3":
            cmd += ["-b:a", f"{int(quality)}k"]
        cmd.append(str(target))
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        if result.returncode != 0:
            return {"success": False, "error": f"ffmpeg hatası: {result.stderr[:300]}"}
        return {"success": True, "message": f"Ses çıkarıldı: {target.name}"}

    # ── HANDLER: Video format dönüştürme ──
    @staticmethod
    def _conv_video(src: Path, target: Path, **kw) -> dict:
        ffmpeg = FileOperationsModule._find_ffmpeg()
        if not ffmpeg:
            return {"success": False, "error": "ffmpeg bulunamadı."}
        import subprocess

        cmd = [
            ffmpeg,
            "-i",
            str(src),
            "-y",
            "-c:v",
            "libx264",
            "-c:a",
            "aac",
            str(target),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
        if result.returncode != 0:
            return {"success": False, "error": f"ffmpeg hatası: {result.stderr[:300]}"}
        return {"success": True, "message": f"Video dönüştürüldü: {target.name}"}

    # ── HANDLER: Video → GIF ──
    @staticmethod
    def _conv_video_to_gif(src: Path, target: Path, **kw) -> dict:
        ffmpeg = FileOperationsModule._find_ffmpeg()
        if not ffmpeg:
            return {"success": False, "error": "ffmpeg bulunamadı."}
        import subprocess

        # İlk 10 saniye, 10fps, 320px genişlik
        cmd = [
            ffmpeg,
            "-i",
            str(src),
            "-y",
            "-t",
            "10",
            "-vf",
            "fps=10,scale=320:-1:flags=lanczos",
            str(target),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            return {"success": False, "error": f"ffmpeg hatası: {result.stderr[:300]}"}
        return {"success": True, "message": f"GIF oluşturuldu: {target.name}"}

    # ── HANDLER: ZIP çıkar ──
    @staticmethod
    def _conv_unzip(src: Path, target: Path, **kw) -> dict:
        import zipfile
        extract_dir = src.parent / src.stem
        extract_dir.mkdir(exist_ok=True)
        with zipfile.ZipFile(str(src), "r") as zf:
            zf.extractall(str(extract_dir))
        return {"success": True, "message": f"Çıkarıldı: {extract_dir.name}/",
                "extract_dir": str(extract_dir)}

    # ── ffmpeg bulma yardımcısı ──
    @staticmethod
    def _find_ffmpeg() -> str:
        """Sistemde ffmpeg'i bul."""
        import shutil as _shutil
        # PATH'te ara
        ff = _shutil.which("ffmpeg")
        if ff:
            return ff
        # Yaygın Windows konumları
        for p in [
            r"C:\ffmpeg\bin\ffmpeg.exe",
            r"C:\Program Files\ffmpeg\bin\ffmpeg.exe",
            r"C:\tools\ffmpeg\bin\ffmpeg.exe",
        ]:
            if Path(p).exists():
                return p
        return ""

    # ================================================================
    # YARDIMCI
    # ================================================================

    @staticmethod
    def _human_size(size_bytes: int) -> str:
        if size_bytes >= 1024**3:
            return f"{size_bytes / 1024**3:.2f} GB"
        elif size_bytes >= 1024**2:
            return f"{size_bytes / 1024**2:.2f} MB"
        elif size_bytes >= 1024:
            return f"{size_bytes / 1024:.1f} KB"
        return f"{size_bytes} B"